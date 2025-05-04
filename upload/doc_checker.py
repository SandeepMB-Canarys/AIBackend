import os
import re
import time
import json
import pandas as pd
import numpy as np
from typing import Optional, List, Dict
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from nltk.tokenize import sent_tokenize
from sklearn.metrics.pairwise import cosine_similarity
from openai import OpenAI, RateLimitError
import anthropic
from sentence_transformers import SentenceTransformer
from django.conf import settings
import nltk
import logging

# ============================== Logging Setup ==============================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)

# ============================== punkt tokenizer==============================

nltk.download('punkt')

# ============================== Token Usage Summary ==============================
token_usage_summary = {
    "prompt_tokens": 0,
    "completion_tokens": 0,
    "latency_seconds": 0.0,
}

def get_token_usage() -> Dict[str, float]:
    return token_usage_summary.copy()

def update_token_usage(prompt_tokens: int, completion_tokens: int, latency: float):
    token_usage_summary["prompt_tokens"] += prompt_tokens
    token_usage_summary["completion_tokens"] += completion_tokens
    token_usage_summary["latency_seconds"] += latency

# ============================== Initialize OpenAI ==============================
def initialize_openai(api_key: str, llm_model: str = 'gpt-4.1-mini'):

    logging.info("Initializing OpenAI client and embeddings.")
    client = OpenAI(api_key=api_key)
    embedding = SentenceTransformer('all-MiniLM-L6-v2')
    logging.info("Clients initialized successfully.")
    return client, embedding

# ============================== Initialize Claude ==============================
def initialize_claude(api_key: str, llm_model: str = 'claude-3'):

    logging.info("Initializing Claude client and embeddings.")
    client = anthropic.Anthropic(api_key=api_key)
    embedding = SentenceTransformer('all-MiniLM-L6-v2')
    logging.info("Clients initialized successfully.")
    return client, embedding

# ============================== Retry Helpers ==============================
def retry_on_rate_limit(func, *args, max_retries=5, initial_wait=1, **kwargs):
    retries = 0
    wait_time = initial_wait
    while retries < max_retries:
        try:
            return func(*args, **kwargs)
        except RateLimitError as e:
            logging.warning(f"Rate limit exceeded. Retrying in {wait_time} seconds...")
            time.sleep(wait_time)
            retries += 1
            wait_time *= 2  # Exponential backoff
    raise Exception(f"Failed after {max_retries} retries.")

# ============================== DOCX Extract ==============================
def extract_preview_text_from_docx(docx_path: str, max_paragraphs: int = 150) -> str:
    logging.info(f"Extracting preview text from {docx_path}.")
    doc = Document(docx_path)
    preview_text = "\n".join(para.text for para in doc.paragraphs[:max_paragraphs] if para.text.strip())
    return preview_text

def extract_full_text_from_docx(docx_path: str) -> str:
    logging.info(f"Extracting full text from {docx_path}.")
    doc = Document(docx_path)
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return full_text

def extract_headings_from_docx(docx_path: str) -> str:
    logging.info(f"Extracting headings from {docx_path}.")
    doc = Document(docx_path)
    headings = [para.text.strip() for para in doc.paragraphs if para.style and para.style.name.startswith("Heading")]
    logging.info(f"Extracted {len(headings)} headings.")
    return "\n".join(headings)

def normalize_doc_type(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"\(.*?\)", "", text)
    text = re.sub(r"s\b", "", text)
    return text.strip()

# ============================== LLM Prompt ==============================
def ask_llm(prompt: str, client, model: str = "gpt-4o", provider: str = "openai", track_usage: bool = True):
    logging.info(f"Sending prompt to {provider} LLM.")
    start_time = time.time()

    try:
        if provider == "openai":
            response = retry_on_rate_limit(
                client.chat.completions.create,
                model=model,
                temperature=0,
                messages=[{"role": "system", "content": prompt}]
            )
            result = response.choices[0].message.content.strip() if response.choices else ""

            # Capture token usage if available
            usage = response.usage if hasattr(response, 'usage') else None
            prompt_tokens = usage.prompt_tokens if usage else 0
            completion_tokens = usage.completion_tokens if usage else 0

        elif provider == "claude":
            response = client.messages.create(
                model=model,
                temperature=0,
                max_tokens=4096,
                system="You are a helpful assistant.",
                messages=[{"role": "user", "content": prompt}]
            )
            result = response.content[0].text.strip() if hasattr(response, 'content') and response.content else ""

            # Claude does not return token usage by default
            prompt_tokens = 0
            completion_tokens = 0

        else:
            raise ValueError(f"Unsupported provider: {provider}")

    except Exception as e:
        logging.error(f"Error in LLM request ({provider}): {e}")
        return "", 0, 0, 0

    latency = time.time() - start_time
    logging.info(f"LLM response received from {provider} in {latency:.2f} seconds.")
    logging.info(f"Prompt tokens: {prompt_tokens}, Completion tokens: {completion_tokens}")

    if track_usage:
        update_token_usage(prompt_tokens, completion_tokens, latency)

    return result, prompt_tokens, completion_tokens, latency

# ============================== Type Detection ==============================
def detect_document_type( docx_path: str, rules_excel_path: str, client, embedding: SentenceTransformer, model: str = "gpt-4", provider: str = "openai") -> str:
    logging.info(f"Detecting document type for {docx_path}.")
    
    # Load and group rule descriptions
    df = pd.read_excel(rules_excel_path, header=1).fillna("")
    grouped = df.groupby("documentType")["checkpointDescription"].apply(lambda x: " ".join(x)).reset_index()

    doc_type_texts = grouped["checkpointDescription"].tolist()
    doc_types = grouped["documentType"].tolist()

    # Extract content from document
    file_name_text = os.path.basename(docx_path).replace("_", " ").replace("-", " ")
    preview_text = extract_preview_text_from_docx(docx_path)
    headings_text = extract_headings_from_docx(docx_path)

    # Generate combined document embedding
    doc_embedding = (
        0.2 * np.array(retry_on_rate_limit(embedding.encode, file_name_text)) +
        0.4 * np.array(retry_on_rate_limit(embedding.encode, preview_text)) +
        0.4 * np.array(retry_on_rate_limit(embedding.encode, headings_text))
    )

    # Compare with known document type embeddings
    type_embeddings = retry_on_rate_limit(embedding.encode, doc_type_texts)
    similarities = cosine_similarity([doc_embedding], type_embeddings)[0]
    best_match_index = np.argmax(similarities)
    best_candidate = doc_types[best_match_index]

    # Prompt construction
    prompt = (
        "You are a document classification assistant.\n\n"
        f"File name: \"{file_name_text}\"\n"
        f"First 2 pages preview:\n{preview_text}\n\n"
        f"Section headings:\n{headings_text}\n\n"
        f"The closest semantic match is \"{best_candidate}\".\n"
        "Suggest the most appropriate document type. Return only the document type."
    )

    # Ask LLM using the selected provider
    result, _, _, _ = ask_llm(prompt, client, model=model, provider=provider)
    doc_type = result.strip()

    logging.info(f"Document type detected: {doc_type}")
    return doc_type

# ============================== Rule Checking ==============================
def check_rules_against_document(docx_path: str, rules_excel_path: str, doc_type: str, client: OpenAI, embedding: SentenceTransformer, model: str = "gpt-4", provider: str = "openai", batch_size: int = 12) -> Optional[pd.DataFrame]:
    logging.info(f"Checking rules for document type: {doc_type} against {docx_path}.")
    df = pd.read_excel(rules_excel_path, header=1).fillna("")
    df["normalized_doc_type"] = df["documentType"].apply(normalize_doc_type)
    norm_detected_type = normalize_doc_type(doc_type)

    df_rules = df[df["normalized_doc_type"] == norm_detected_type]
    if df_rules.empty:
        logging.warning(f"No rules found for document type: {doc_type}")
        return None

    document_text = extract_full_text_from_docx(docx_path)
    results_list = []
    rule_batches = [df_rules.iloc[i:i + batch_size] for i in range(0, len(df_rules), batch_size)]

    def process_batch(batch_df):
        rules_text = "\n".join([ 
            f"{idx + 1}. Rule: {row['ruleType']} | Subrule: {row['checkpointDescription']}"
            for idx, row in batch_df.iterrows()
        ]) 

        prompt = (
            "You are a document compliance checker. For each rule and subrule below, check if the document content meets the requirements.\n\n"
            "For each rule, respond in the format:\n"
            "1. ✅ Rule if met | Subrule: subrule text | Reason: reason\n"
            "2. ❌ Rule if not met | Subrule: subrule text | Reason: reason\n"
            f"\nDocument Content:\n{document_text}\n\nRules and Subrules:\n{rules_text}\n"
        )
        response = ask_llm(prompt, client, model=model, provider=provider)

        # handle tuple
        response_text = response[0] if isinstance(response, tuple) else response

        response_text = response_text.replace("\\n", "\n")

        for line in response_text.splitlines():
            print("line", line)

            match = re.match(r"\d+\.\s*(✅|❌)\s*(.*?)\s*\|\s*Subrule:\s*(.*?)\s*\|\s*Reason:\s*(.*)", line.strip())
            print("match", match)

            if match:
                status, rule, subrule, reason = match.groups()
                            
                results_list.append({
                    "RuleType": rule.strip(),
                    "CheckpointDescription": subrule.strip(),
                    "Status": "✅ Rule Met" if status == "✅" else "❌ Not Met",
                    "Reason": reason.strip()
                })

        print("results_list", results_list)

    with ThreadPoolExecutor() as executor:
        executor.map(process_batch, rule_batches)

    results_df = pd.DataFrame(results_list)
    print("results_df", results_df)

    # Store the results in an Excel file after processing
    if results_df.empty:
        logging.warning("No results to store.")
    else:
        logging.info("Storing report...")
        output_dir = os.path.join(settings.MEDIA_ROOT, 'report')
        
        # Store report and get updated DataFrame with report_path
        results_df = store_report(results_df, doc_name=os.path.basename(docx_path), output_dir=output_dir)

    return results_df

# ============================== Store Report ==============================
def store_report(results_df: pd.DataFrame, doc_name: str, output_dir: str) -> pd.DataFrame:
    if not results_df.empty:
        # Get the current date and time
        current_datetime = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        
        # Format the filename to include date, time, and document name
        file_name = f"{doc_name}_{current_datetime}_report.xlsx"
        file_path = os.path.join(output_dir, file_name)
        
        # Create the output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Write the DataFrame to an Excel file
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            results_df.to_excel(writer, index=False, sheet_name="Compliance Report")
            
        logging.info(f"Report stored successfully at {file_path}.")
        
        # Set the report path attribute on the DataFrame
        results_df.report_path = file_path
    else:
        logging.warning("No results to store.")
        results_df.report_path = None

    return results_df

